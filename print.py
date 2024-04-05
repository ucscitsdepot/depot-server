import pypandoc
from docx import Document
import sys
import cups
import shutil
import mammoth
from html2image import Html2Image

from flask import Flask, render_template, request
import os
from datetime import date
hti = Html2Image(custom_flags=['--no-sandbox'], size=(800, 1000))
rtf_path = '/home/depot/Auto-Label-Generator/ITS-Shipping-Form.rtf'
docx_path = '/home/depot/Auto-Label-Generator/ITS-Shipping-Form.docx'
new_docx_path = '/home/depot/Auto-Label-Generator/ITS-Shipping-Form-Copy.docx'
shutil.copy(docx_path, new_docx_path)
printer_name = "printername"
cmd = "lp -o fill /home/depot/Auto-Label-Generator/blue_page.png"
current_date = date.today().strftime("%m/%d/%Y")






def convert_rtf_to_docx(rtf_path, docx_path):
    output = pypandoc.convert_file(rtf_path, 'docx', outputfile=docx_path)

def replace_string_in_docx(docx_path, old_string, new_string):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_string in paragraph.text:
            paragraph.text = paragraph.text.replace(old_string, new_string)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_string in cell.text:
                    cell.text = cell.text.replace(old_string, new_string)
    doc.save(docx_path)
                # do something with the cell
                

def convert_docx_to_rtf(docx_path, rtf_path):
    output = pypandoc.convert_file(docx_path, 'rtf', outputfile=rtf_path)

def print_file_using_pycups(file_path, printer_name=None):
    conn = cups.Connection()
    
    # Get the default printer if not specified
    if not printer_name:
        printer_name = conn.getDefault()
    
    # Print the file
    job_id = conn.printFile(printer_name, file_path, "Print Job", {})
    print(f'Successfully sent {file_path} to the printer (Job ID: {job_id}).')

def adjust_string_length(variable, length):
    underscores = '_' * (length - len(variable))
    return variable + underscores


def printcall(name, phone, address1, address2, city, state, zip_code, mailcode, tracking_email, approver, ritm):
    replace_string_in_docx(new_docx_path, 'Name ____________________________', 'Name: %s' % adjust_string_length(name, 29))
    replace_string_in_docx(new_docx_path, 'Date _____________', 'Date: %s' % adjust_string_length(current_date, 10))
    replace_string_in_docx(new_docx_path, 'Phone _____________________', 'Phone: %s' % adjust_string_length(phone, 5))
    replace_string_in_docx(new_docx_path, 'Address Line 1 __________________________________________________________________', 'Address1: %s' % adjust_string_length(address1, 5))
    replace_string_in_docx(new_docx_path, 'Address Line 2 __________________________________________________________________', 'Address2: %s' % adjust_string_length(address2, 5))
    replace_string_in_docx(new_docx_path, 'City ____________________', 'City: %s' % adjust_string_length(city, 7))
    replace_string_in_docx(new_docx_path, 'State _________', 'State: %s' % adjust_string_length(state, 2))
    replace_string_in_docx(new_docx_path, 'ZIP _____________', 'Zip: %s' % adjust_string_length(zip_code, 5))
    replace_string_in_docx(new_docx_path, 'MailCode ________', 'Mailcode: %s' % adjust_string_length(mailcode, 2))
    replace_string_in_docx(new_docx_path, 'depot@ucsc.edu | __________________________________', 'depot@ucsc.edu | %s' % adjust_string_length(tracking_email, 14))
    replace_string_in_docx(new_docx_path, 'MailCode Approver _____________________________', 'MailCode Approver: %s' % adjust_string_length(approver, 8))
    replace_string_in_docx(new_docx_path, 'RITM00_____________', '%s' % adjust_string_length(ritm, 5))
    # replace_string_in_docx(new_docx_path, 'INC0_____________', '%s' % adjust_string_length(inc, 5))
    
    custom_styles = "b => i"
    with open(new_docx_path, "rb") as docx_file:
        
        result = mammoth.convert_to_html(docx_file, style_map=custom_styles)
        text = result.value
        with open('/home/depot/Auto-Label-Generator/output.html', 'w') as html_file:
            html_file.write(text)
    hti.screenshot(html_file='/home/depot/Auto-Label-Generator/output.html', save_as='blue_page.png')
    os.system(cmd)




# custom_styles = "b => i"
# with open(new_docx_path, "rb") as docx_file:
        
#         result = mammoth.convert_to_html(docx_file, style_map=custom_styles)
#         text = result.value
#         with open('/home/depot/Auto-Label-Generator/output.html', 'w') as html_file:
#             html_file.write(text)
# hti.screenshot(html_file='/home/depot/Auto-Label-Generator/output.html', save_as='blue_page.png')
    
# os.system(cmd)
        






